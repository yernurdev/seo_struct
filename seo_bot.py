import os
import io
import json
import logging
import asyncio
import re
import aiohttp
from dotenv import load_dotenv

import pandas as pd
from docx import Document

from aiogram import Bot, Dispatcher, F
from aiogram.types import Message, CallbackQuery, InlineKeyboardMarkup, InlineKeyboardButton, BufferedInputFile
from aiogram.filters import CommandStart, Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

from storage_db import StorageDB

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(name)s - %(message)s",
    handlers=[
        logging.FileHandler("bot_seo.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# --- Внедрение эмбеддингов ---
try:
    from sentence_transformers import SentenceTransformer, util
    logger.info("Загрузка модели sentence-transformers...")
    # Using all-MiniLM-L6-v2 as requested/standard
    embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
    logger.info("Модель успешно загружена.")
except ImportError:
    logger.warning("sentence-transformers не установлена. Оценка плотности и скоринг будут недоступны.")
    embedding_model = None

# Загрузка переменных окружения
load_dotenv()
TG_TOKEN = os.getenv("TG_TOKEN")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
PIXEL_TOOLS_KEY = os.getenv("PIXEL_TOOLS_KEY")
USE_MOCK = os.getenv("USE_MOCK", "0") == "1"

if not TG_TOKEN or not OPENROUTER_API_KEY:
    logger.error("TG_TOKEN или OPENROUTER_API_KEY не найдены. Установите их в .env файле.")
    exit(1)

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

class Response:
    def __init__(self, text):
        self.text = text

class OpenRouterChat:
    def __init__(self):
        self.history = []

    async def send_message(self, content: str) -> Response:
        self.history.append({"role": "user", "content": content})

        if USE_MOCK:
            logger.info("Используется MOCK для OpenRouter.")
            await asyncio.sleep(1)
            
            # Эмуляция ответов
            if "составь SEO-структуру" in content:
                ans = "Структура статьи:\n\nX1. Введение\nX2. Основная часть\n- Блок А\n- Блок Б\nX3. Заключение"
            elif "Техническое Задание" in content:
                ans = "ТЗ:\nLLSI слова: пример, тема, тест, текст, структура\nОбъем: 2000 сим."
            elif "Выпиши как JSON массив" in content:
                ans = '["X1. Введение", "X2. Основная часть", "X3. Заключение"]'
            elif "Напиши текст строго для блока:" in content:
                ans = "Это тестовый текст для одного из блоков. В нем есть много полезной информации по теме, также пример и тест."
            else:
                ans = "Mock response"
            
            self.history.append({"role": "assistant", "content": ans})
            return Response(ans)

        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": "google/gemini-2.5-flash",
            "max_tokens": 4096,
            "messages": self.history,
        }

        async with aiohttp.ClientSession() as session:
            async with session.post(OPENROUTER_URL, headers=headers, json=payload) as resp:
                data = await resp.json()
                if "error" in data:
                    error_msg = data["error"].get("message", str(data["error"]))
                    raise Exception(f"OpenRouter API error: {error_msg}")
                assistant_msg = data["choices"][0]["message"]["content"]
                self.history.append({"role": "assistant", "content": assistant_msg})
                return Response(assistant_msg)

bot = Bot(token=TG_TOKEN)
dp = Dispatcher()

DB = StorageDB()

class SeoProcess(StatesGroup):
    waiting_for_file = State()
    step0_structure = State()
    step1_tz = State()
    step2_text = State()

USER_CHATS = {}
USER_DATA = {}

def get_or_create_chat(user_id: int):
    if user_id not in USER_CHATS:
        USER_CHATS[user_id] = OpenRouterChat()
    return USER_CHATS[user_id]

def clear_user_session(user_id: int):
    if user_id in USER_CHATS:
        del USER_CHATS[user_id]
    if user_id in USER_DATA:
        del USER_DATA[user_id]

async def safe_llm_request(chat, prompt, max_retries=3, delay=5):
    for attempt in range(max_retries):
        try:
            return await chat.send_message(prompt)
        except Exception as e:
            if ("503" in str(e) or "429" in str(e) or "rate" in str(e).lower()) and attempt < max_retries - 1:
                logger.warning(f"Ошибка LLM ({attempt + 1}/{max_retries}): {e}. Ждем {delay} сек.")
                await asyncio.sleep(delay)
                continue
            raise e

async def send_long_text(status_msg: Message, text: str, reply_markup=None):
    max_len = 4000
    parts = [text[i:i+max_len] for i in range(0, len(text), max_len)]
    for i, part in enumerate(parts):
        markup = reply_markup if i == len(parts) - 1 else None
        if i == 0:
            await status_msg.edit_text(part, reply_markup=markup)
        else:
            await status_msg.answer(part, reply_markup=markup)

async def get_top_competitors(keys):
    if USE_MOCK:
        logger.info("MOCK: возвращаем тестовые URL конкурентов.")
        await asyncio.sleep(1)
        return ["https://mock-competitor.com/1", "https://mock-competitor.com/2"]

    if not PIXEL_TOOLS_KEY:
        logger.error("PIXEL_TOOLS_KEY не установлен.")
        return []

    API_URL = "https://tools.pixelplus.ru/api/top10"
    REGION = "213"
    SEARCH_ENGINE = "yandex"
    
    async with aiohttp.ClientSession() as session:
        payload = {'queries': json.dumps(keys), 'lr': REGION, 'ss': SEARCH_ENGINE, 'deep': 10}
        params = {'key': PIXEL_TOOLS_KEY}
        try:
            async with session.post(API_URL, params=params, data=payload) as resp:
                result = await resp.json()
                if 'report_id' not in result:
                    return []
                report_id = result['report_id']
        except Exception as e:
            logger.error(f"Ошибка PixelPlus (постановка): {e}")
            return []

        attempt = 0
        while attempt < 30:
            attempt += 1
            await asyncio.sleep(10)
            try:
                params = {'key': PIXEL_TOOLS_KEY, 'report_id': report_id}
                async with session.get(API_URL, params=params) as resp:
                    data = await resp.json()
                    if data.get('status') == 'success' and data.get('msg') == 'completed':
                        urls_found = set()
                        results_data = data.get('response', {}).get('response', {}).get('result', {})
                        for query_val in results_data.values():
                            platform_key = 'yandex_urls' if SEARCH_ENGINE.startswith('yandex') else 'google_urls'
                            raw_urls = query_val.get(platform_key, [])
                            for item in raw_urls:
                                if 'url' in item:
                                    urls_found.add(item['url'])
                        return list(urls_found)[:20]
                    elif data.get('code') == -50 or data.get('error') == 'In progress':
                        continue
                    else:
                        break
            except Exception as e:
                logger.error(f"Ошибка PixelPlus (результат): {e}")
                break
        return []

def calculate_block_density(text: str, lsi_words: list):
    """Оценка плотности LSI слов через эмбеддинги (семантическое сходство) и физическое вхождение."""
    if not lsi_words:
        return 0.0
    
    # 1. Считаем физическую плотность
    words_in_text = text.lower().split()
    total_words = len(words_in_text)
    if total_words == 0: return 0.0

    found_count = sum(1 for w in lsi_words if w.lower() in text.lower())
    density_percent = (found_count / max(1, len(lsi_words))) * 100

    # 2. Оценка структуры/семантическая плотность (если загружена модель)
    if embedding_model:
        lsi_text = " ".join(lsi_words)
        lsi_emb = embedding_model.encode(lsi_text)
        text_emb = embedding_model.encode(text)
        similarity = util.cos_sim(lsi_emb, text_emb).item() * 100
        # Итоговая комбинированная плотность по блокам - микс прямого вхождения и семантики
        return (density_percent + max(0, similarity)) / 2

    return density_percent

def extract_lsi_words(tz_text: str):
    """Извлекает SI-слова из сгенерированного ТЗ, анализируя текст."""
    match = re.search(r'LLSI слова:(.*?)(?:\n|$)', tz_text, re.IGNORECASE)
    if match:
        words = match.group(1).split(',')
        return [w.strip() for w in words if w.strip()]
    return ["анализ", "качество", "услуга"] # Fallback

@dp.message(CommandStart())
async def cmd_start(message: Message, state: FSMContext):
    await state.set_state(SeoProcess.waiting_for_file)
    clear_user_session(message.from_user.id)
    await message.answer(
        "👋 Привет! Отправьте мне Excel файл с ключами (столбцы 'Фраза' и 'WS').",
        parse_mode="HTML"
    )

@dp.message(Command("cancel"))
async def cmd_cancel(message: Message, state: FSMContext):
    await state.clear()
    clear_user_session(message.from_user.id)
    await message.answer("Отменено.")

@dp.message(SeoProcess.waiting_for_file, F.document)
async def process_excel_file(message: Message, state: FSMContext):
    document = message.document
    if not document.file_name.lower().endswith(('.xlsx', '.csv')):
        return await message.answer("❌ Нужен .xlsx или .csv!")

    status_msg = await message.answer("📥 Скачиваю файл...")
    user_id = message.from_user.id

    try:
        file_in_memory = io.BytesIO()
        await bot.download(document, destination=file_in_memory)
        file_in_memory.seek(0)
        
        df = pd.read_excel(file_in_memory) if document.file_name.lower().endswith('.xlsx') else pd.read_csv(file_in_memory)
        if 'Фраза' not in df.columns or 'WS' not in df.columns:
            return await status_msg.edit_text("❌ Ошибка: нет столбцов 'Фраза', 'WS'.")

        keys = df['Фраза'].astype(str).tolist()
        ws_values = df['WS'].tolist()
        
        project_id = str(message.from_user.id)

        DB.add_keywords(project_id, keys, ws_values)
        USER_DATA[user_id] = {
        'keys': keys,
        'project_id': str(user_id)
        }


        await status_msg.edit_text("🔍 Парсим конкурентов...")
        top_urls = await get_top_competitors(keys)

        DB.add_competitors(project_id, top_urls)

        

        await status_msg.edit_text("🧠 Генерирую структуру...")
        keys_str = "\n".join([f"- {k} ({w})" for k, w in zip(keys, ws_values)])
        prompt = f"Рассчитай оптимальный объем текста и составь SEO-структуру (H1, H2).\nКлючи:\n{keys_str}"
        if top_urls:
            prompt += f"\nURL конкурентов:\n" + "\n".join(top_urls[:5])

        chat = get_or_create_chat(user_id)
        response = await safe_llm_request(chat, prompt)
        
        USER_DATA[user_id]['structure'] = response.text
        
        markup = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="Дальше", callback_data="step0_continue")]
        ])
        await send_long_text(status_msg, response.text, markup)
        await state.set_state(SeoProcess.step0_structure)

    except Exception as e:
        logger.error(f"Ошибка {e}")
        await status_msg.edit_text("❌ Ошибка.")

@dp.callback_query(SeoProcess.step0_structure, F.data == "step0_continue")
async def step0_callback(call: CallbackQuery, state: FSMContext):
    user_id = call.from_user.id
    chat = USER_CHATS.get(user_id)
    if not chat: return await call.message.edit_text("❌ Сессия устарела.")
    await call.message.edit_reply_markup(reply_markup=None)
    
    status_msg = await call.message.answer("📝 Формирую ТЗ...")
    try:
        prompt = "Составь ТЗ по шаблону:\n- Требования к блокам\n- LLSI слова\n- Объем"
        response = await safe_llm_request(chat, prompt)
        USER_DATA[user_id]['tz'] = response.text
        markup = InlineKeyboardMarkup(inline_keyboard=[[InlineKeyboardButton(text="Дальше", callback_data="step1_continue")]])
        await send_long_text(status_msg, response.text, markup)
        await state.set_state(SeoProcess.step1_tz)
    except Exception as e:
        await status_msg.edit_text("❌ Ошибка.")

@dp.callback_query(SeoProcess.step1_tz, F.data == "step1_continue")
async def step1_callback(call: CallbackQuery, state: FSMContext):
    user_id = call.from_user.id
    chat = USER_CHATS.get(user_id)

    project_id = USER_DATA.get(user_id, {}).get('project_id')

    if not chat: 
        return await call.message.edit_text("❌ Сессия устарела.")
    
    await call.message.edit_reply_markup(reply_markup=None)

    status_msg = await call.message.answer("✍️ Генерирую текст по блокам... ⏳")
    try:
        # Извлекаем список блоков из ТЗ
        prompt_blocks = "Выпиши как JSON массив список всех подзаголовков (блоков) из ТЗ, для которых нужно написать текст. Просто массив строк, без markdown."
        blocks_resp = await safe_llm_request(chat, prompt_blocks)
        blocks_text = blocks_resp.text.replace("```json", "").replace("```", "").strip()
        try:
            blocks = json.loads(blocks_text)
        except:
            blocks = ["Один большой блок (ошибка парсинга)"]

        lsi_words = extract_lsi_words(USER_DATA.get(user_id, {}).get('tz', ''))
        USER_DATA[user_id]['analytics'] = []
        full_text = ""

        # Поблочная генерация с подсчетом плотности SI
        for i, block_name in enumerate(blocks):
            await status_msg.edit_text(f"✍️ Генерация: блок {i+1} из {len(blocks)} ({block_name})... ⏳")
            block_prompt = f"Напиши текст строго для блока: '{block_name}'. Учитывай наше ТЗ. Органично используй эти LLSI слова: {', '.join(lsi_words)}."
            block_resp = await safe_llm_request(chat, block_prompt)

            DB.add_text(project_id, block_name, block_resp.text)
            
            density = calculate_block_density(block_resp.text, lsi_words)
            USER_DATA[user_id]['analytics'].append(f"Блок '{block_name}': Плотность LLSI = {density:.1f}%")
            
            full_text += f"\n\n### {block_name}\n\n{block_resp.text}"

        USER_DATA[user_id]['text'] = full_text
        
        markup = InlineKeyboardMarkup(inline_keyboard=[[InlineKeyboardButton(text="Сформировать Word-файл", callback_data="step2_continue")]])
        await send_long_text(status_msg, full_text, markup)
        await state.set_state(SeoProcess.step2_text)

    except Exception as e:
        logger.error(f"Ошибка генерации текста: {e}")
        await status_msg.edit_text("❌ Ошибка генерации текста.")

@dp.callback_query(SeoProcess.step2_text, F.data == "step2_continue")
async def step2_callback(call: CallbackQuery, state: FSMContext):
    user_id = call.from_user.id
    if not USER_CHATS.get(user_id): return await call.message.edit_text("❌ Сессия устарела.")
    await call.message.edit_reply_markup(reply_markup=None)

    status_msg = await call.message.answer("📄 Формирую Word-документ (с отчетом по SI)...")
    try:
        data = USER_DATA.get(user_id, {})
        tz_content = data.get('tz', '')
        text_content = data.get('text', '')
        analytics = data.get('analytics', [])
        
        doc = Document()
        doc.add_heading("Аналитика и скоринг", level=1)
        doc.add_paragraph("Оценка плотности LSI слов (комбинация частотности и семантики):")
        for an in analytics:
            doc.add_paragraph(an)

        doc.add_page_break()
        doc.add_heading("Техническое Задание (ТЗ)", level=1)
        doc.add_paragraph(tz_content)
        
        doc.add_page_break()
        doc.add_heading("SEO-текст", level=1)
        doc.add_paragraph(text_content)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        document = BufferedInputFile(file_stream.read(), filename="seo_article.docx")
        await call.message.answer_document(document=document, caption="✅ Статья и отчет сформированы!\n\n" + "\n".join(analytics))
        await status_msg.delete()
        
        await state.clear()
        clear_user_session(user_id)
    except Exception as e:
        logger.error(f"Ошибка Word: {e}")
        await status_msg.edit_text("❌ Ошибка Word.")

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("Бот остановлен.")

